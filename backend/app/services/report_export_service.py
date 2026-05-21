from __future__ import annotations

import json
import re
from datetime import datetime
from dataclasses import dataclass, field
from html import escape
from pathlib import Path
from typing import Literal

from app.core.exceptions import ArtifactNotFoundError, DependencyUnavailableError, InputValidationError
from app.core.settings import Settings

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ModuleNotFoundError as exc:  # pragma: no cover - 依赖缺失分支通过路由测试覆盖
    Document = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None
    RGBColor = None
    DOCX_IMPORT_ERROR: ModuleNotFoundError | None = exc
else:
    DOCX_IMPORT_ERROR = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        HRFlowable,
        NextPageTemplate,
        PageBreak,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.platypus.tableofcontents import TableOfContents
except ModuleNotFoundError as exc:  # pragma: no cover - 依赖缺失分支通过路由测试覆盖
    colors = None
    A4 = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    cm = None
    pdfmetrics = None
    UnicodeCIDFont = None
    BaseDocTemplate = None
    Frame = None
    HRFlowable = None
    NextPageTemplate = None
    PageBreak = None
    PageTemplate = None
    Paragraph = None
    Spacer = None
    Table = None
    TableStyle = None
    TableOfContents = None
    PDF_IMPORT_ERROR: ModuleNotFoundError | None = exc
else:
    PDF_IMPORT_ERROR = None

ReportExportFormat = Literal["md", "docx", "pdf"]
PdfCoverDateMode = Literal["today", "custom", "hide"]

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*:?-+:?\s*\|?\s*$")
HORIZONTAL_RULE_RE = re.compile(r"^\s*([-*_])(?:\s*\1){2,}\s*$")
INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")

EXPORT_FILENAMES: dict[ReportExportFormat, str] = {
    "md": "report.md",
    "docx": "report.docx",
    "pdf": "report.pdf",
}

EXPORT_MEDIA_TYPES: dict[ReportExportFormat, str] = {
    "md": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

PDF_COVER_TITLE = "道路交通事故分析报告"
PDF_COVER_SUBTITLE = "事故事实梳理、责任分析与研判文书"
PDF_COMPILED_BY = "锐鉴安途道路交通事故分析系统"
PDF_GENERIC_BODY_TITLES = {"交通事故分析报告", "道路交通事故分析报告"}
PDF_COLOR_SURFACE = "#F5F1EA"
PDF_COLOR_PANEL = "#E8DED0"
PDF_COLOR_ACCENT = "#A97B47"
PDF_COLOR_INK = "#1E293B"
PDF_COLOR_MUTED = "#6B7280"
PDF_COLOR_LINE = "#D7CBB8"
PDF_COLOR_HEADER = "#334155"


@dataclass
class InlineRun:
    text: str
    bold: bool = False
    code: bool = False


@dataclass
class MarkdownBlock:
    kind: Literal["heading", "paragraph", "list", "table", "hr"]
    level: int = 0
    ordered: bool = False
    runs: list[InlineRun] = field(default_factory=list)
    items: list[list[InlineRun]] = field(default_factory=list)
    rows: list[list[list[InlineRun]]] = field(default_factory=list)


@dataclass(frozen=True)
class PdfCoverOptions:
    title: str | None = None
    subtitle: str | None = None
    compiled_by: str | None = None
    date_mode: PdfCoverDateMode = "today"
    date_text: str | None = None


@dataclass(frozen=True)
class ResolvedPdfCoverConfig:
    title: str
    subtitle: str
    compiled_by: str
    date_text: str | None
    report_number: str

    def to_cache_payload(self) -> dict[str, str | None]:
        return {
            "title": self.title,
            "subtitle": self.subtitle,
            "compiled_by": self.compiled_by,
            "date_text": self.date_text,
            "report_number": self.report_number,
        }


class _StyledPdfDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, title: str, author: str, subject: str):
        super().__init__(
            filename,
            pagesize=A4,
            topMargin=2.2 * cm,
            bottomMargin=2.0 * cm,
            leftMargin=2.0 * cm,
            rightMargin=2.0 * cm,
            title=title,
            author=author,
            subject=subject,
        )
        self.report_title = title
        self.report_author = author
        self._outline_counter = 0

        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="content")
        self.addPageTemplates(
            [
                PageTemplate(id="Cover", frames=[frame], onPage=self._draw_cover_page),
                PageTemplate(id="Body", frames=[frame], onPage=self._draw_body_page),
            ]
        )

    def beforeDocument(self) -> None:
        self._outline_counter = 0

    def _draw_cover_page(self, canvas, doc) -> None:  # noqa: ANN001
        page_width, page_height = doc.pagesize
        canvas.saveState()
        canvas.setFillColor(colors.HexColor(PDF_COLOR_SURFACE))
        canvas.rect(0, 0, page_width, page_height, stroke=0, fill=1)
        canvas.setStrokeColor(colors.HexColor(PDF_COLOR_LINE))
        canvas.setLineWidth(0.75)
        canvas.rect(
            doc.leftMargin - 0.35 * cm,
            doc.bottomMargin - 0.2 * cm,
            doc.width + 0.7 * cm,
            doc.height + 0.4 * cm,
            stroke=1,
            fill=0,
        )

        canvas.setFillColor(colors.HexColor(PDF_COLOR_PANEL))
        canvas.circle(page_width - doc.rightMargin - 0.9 * cm, page_height - doc.topMargin - 0.9 * cm, 0.62 * cm, stroke=0, fill=1)
        canvas.circle(doc.leftMargin + 0.1 * cm, doc.bottomMargin + 0.5 * cm, 0.22 * cm, stroke=0, fill=1)

        canvas.setStrokeColor(colors.HexColor(PDF_COLOR_ACCENT))
        canvas.setLineWidth(1.15)
        canvas.line(doc.leftMargin, page_height - doc.topMargin - 0.85 * cm, doc.leftMargin + 3.6 * cm, page_height - doc.topMargin - 0.85 * cm)

        canvas.setStrokeColor(colors.HexColor(PDF_COLOR_LINE))
        canvas.setLineWidth(0.7)
        canvas.line(doc.leftMargin, doc.bottomMargin + 1.5 * cm, page_width - doc.rightMargin, doc.bottomMargin + 1.5 * cm)

        canvas.setFont("STSong-Light", 8.8)
        canvas.setFillColor(colors.HexColor(PDF_COLOR_MUTED))
        canvas.drawString(doc.leftMargin, page_height - doc.topMargin - 0.45 * cm, "锐鉴安途事故分析文书")
        canvas.drawRightString(page_width - doc.rightMargin, doc.bottomMargin + 0.92 * cm, self.report_author)
        canvas.restoreState()

    def _draw_body_page(self, canvas, doc) -> None:  # noqa: ANN001
        page_width, page_height = doc.pagesize
        display_page = max(doc.page - 1, 1)

        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(PDF_COLOR_LINE))
        canvas.setLineWidth(0.6)
        canvas.line(doc.leftMargin, page_height - 1.25 * cm, page_width - doc.rightMargin, page_height - 1.25 * cm)
        canvas.line(doc.leftMargin, 1.15 * cm, page_width - doc.rightMargin, 1.15 * cm)

        canvas.setFont("STSong-Light", 8.8)
        canvas.setFillColor(colors.HexColor(PDF_COLOR_MUTED))
        canvas.drawString(doc.leftMargin, page_height - 0.92 * cm, self.report_author)
        canvas.drawRightString(page_width - doc.rightMargin, page_height - 0.92 * cm, self.report_title)

        canvas.setFillColor(colors.HexColor(PDF_COLOR_MUTED))
        canvas.drawString(doc.leftMargin, 0.72 * cm, "交通事故责任分析报告")
        canvas.drawRightString(page_width - doc.rightMargin, 0.72 * cm, f"第 {display_page} 页")
        canvas.restoreState()

    def afterFlowable(self, flowable) -> None:  # noqa: ANN001
        toc_level = getattr(flowable, "_toc_level", None)
        heading_text = str(getattr(flowable, "_plain_text", "") or "").strip()
        if toc_level is None or not heading_text:
            return

        bookmark = f"report-heading-{self._outline_counter}"
        self._outline_counter += 1
        self.canv.bookmarkPage(bookmark)
        outline_level = max(int(toc_level), 0)
        try:
            self.canv.addOutlineEntry(heading_text, bookmark, outline_level, False)
        except ValueError:
            self.canv.addOutlineEntry(heading_text, bookmark, 0, False)
        self.notify("TOCEntry", (int(toc_level), heading_text, max(self.page - 1, 1), bookmark))


class ReportExportService:
    _pdf_font_registered = False

    def __init__(self, settings: Settings):
        self.settings = settings

    def get_export_path(
        self,
        trace_id: str,
        export_format: ReportExportFormat,
        pdf_cover_options: PdfCoverOptions | None = None,
    ) -> Path:
        output_dir = self._resolve_output_dir(trace_id)
        report_markdown_path, report_markdown = self._ensure_report_markdown(output_dir)

        if export_format == "md":
            return report_markdown_path

        target_path = output_dir / EXPORT_FILENAMES[export_format]
        template_mtime_ns = Path(__file__).stat().st_mtime_ns
        blocks = self._parse_markdown_blocks(report_markdown)
        source_mtime_ns = report_markdown_path.stat().st_mtime_ns

        if export_format == "pdf":
            cover_config = self._resolve_pdf_cover_config(
                body_blocks=self._strip_pdf_cover_title_block(blocks),
                trace_id=trace_id,
                options=pdf_cover_options,
            )
            if self._is_pdf_export_current(target_path, source_mtime_ns, template_mtime_ns, cover_config):
                return target_path
            self._build_pdf(target_path, blocks, trace_id, cover_config)
            self._write_pdf_export_meta(target_path, source_mtime_ns, template_mtime_ns, cover_config)
            return target_path

        if target_path.exists() and target_path.stat().st_mtime_ns >= max(
            source_mtime_ns,
            template_mtime_ns,
        ):
            return target_path

        if export_format == "docx":
            self._build_docx(target_path, blocks, report_markdown, trace_id)
            return target_path

        raise InputValidationError(f"暂不支持的导出格式: {export_format}")

    def build_download_name(self, trace_id: str, export_format: ReportExportFormat) -> str:
        return f"traffic-accident-report-{trace_id}.{export_format}"

    @staticmethod
    def get_media_type(export_format: ReportExportFormat) -> str:
        return EXPORT_MEDIA_TYPES[export_format]

    def _resolve_output_dir(self, trace_id: str) -> Path:
        cleaned_trace_id = str(trace_id or "").strip()
        if not cleaned_trace_id:
            raise InputValidationError("报告 trace_id 不能为空。")

        output_dir = (self.settings.output_dir_path / cleaned_trace_id).resolve()
        if not output_dir.is_relative_to(self.settings.output_dir_path):
            raise InputValidationError("报告输出目录非法。")
        if not output_dir.exists() or not output_dir.is_dir():
            raise ArtifactNotFoundError(f"报告输出目录不存在: {cleaned_trace_id}")
        return output_dir

    def _ensure_report_markdown(self, output_dir: Path) -> tuple[Path, str]:
        report_markdown_path = output_dir / EXPORT_FILENAMES["md"]
        if report_markdown_path.exists():
            content = report_markdown_path.read_text(encoding="utf-8").strip()
            if content:
                return report_markdown_path, content

        report_json_path = output_dir / "report.json"
        if not report_json_path.exists():
            raise ArtifactNotFoundError("报告 Markdown 不存在，且无法从 report.json 恢复。")

        try:
            payload = json.loads(report_json_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            raise InputValidationError("报告 JSON 损坏，无法生成导出文件。") from exc

        report_markdown = str(payload.get("report_markdown") or "").strip()
        if not report_markdown:
            raise InputValidationError("报告正文为空，无法导出。")

        report_markdown_path.write_text(report_markdown, encoding="utf-8")
        return report_markdown_path, report_markdown

    def _parse_markdown_blocks(self, markdown_text: str) -> list[MarkdownBlock]:
        lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        blocks: list[MarkdownBlock] = []
        paragraph_lines: list[str] = []
        line_index = 0

        def flush_paragraph() -> None:
            if not paragraph_lines:
                return
            text = "\n".join(line.strip() for line in paragraph_lines).strip()
            if text:
                blocks.append(MarkdownBlock(kind="paragraph", runs=self._parse_inline_runs(text)))
            paragraph_lines.clear()

        while line_index < len(lines):
            raw_line = lines[line_index]
            stripped = raw_line.strip()

            if not stripped:
                flush_paragraph()
                line_index += 1
                continue

            heading_match = HEADING_RE.match(stripped)
            if heading_match:
                flush_paragraph()
                blocks.append(
                    MarkdownBlock(
                        kind="heading",
                        level=len(heading_match.group(1)),
                        runs=self._parse_inline_runs(heading_match.group(2).strip()),
                    )
                )
                line_index += 1
                continue

            if HORIZONTAL_RULE_RE.match(stripped):
                flush_paragraph()
                blocks.append(MarkdownBlock(kind="hr"))
                line_index += 1
                continue

            if stripped.startswith("|"):
                table_lines: list[str] = []
                cursor = line_index
                while cursor < len(lines) and lines[cursor].strip().startswith("|"):
                    table_lines.append(lines[cursor].strip())
                    cursor += 1
                if len(table_lines) >= 2 and TABLE_SEPARATOR_RE.match(table_lines[1]):
                    flush_paragraph()
                    rows: list[list[list[InlineRun]]] = []
                    for table_line in (table_lines[0], *table_lines[2:]):
                        cells = [self._parse_inline_runs(cell) for cell in self._split_table_row(table_line)]
                        if cells:
                            rows.append(cells)
                    if rows:
                        blocks.append(MarkdownBlock(kind="table", rows=rows))
                        line_index = cursor
                        continue

            bullet_match = BULLET_RE.match(stripped)
            ordered_match = ORDERED_RE.match(stripped)
            if bullet_match or ordered_match:
                flush_paragraph()
                items: list[list[InlineRun]] = []
                ordered = ordered_match is not None
                cursor = line_index
                while cursor < len(lines):
                    current = lines[cursor].strip()
                    current_match = ORDERED_RE.match(current) if ordered else BULLET_RE.match(current)
                    if current_match:
                        item_lines = [current_match.group(1).strip()]
                        cursor += 1
                        while cursor < len(lines):
                            continuation = lines[cursor]
                            continuation_stripped = continuation.strip()
                            if not continuation_stripped:
                                cursor += 1
                                break
                            if self._looks_like_new_block(continuation_stripped):
                                break
                            item_lines.append(continuation_stripped)
                            cursor += 1
                        items.append(self._parse_inline_runs("\n".join(item_lines).strip()))
                        continue
                    break
                if items:
                    blocks.append(MarkdownBlock(kind="list", ordered=ordered, items=items))
                    line_index = cursor
                    continue

            paragraph_lines.append(raw_line)
            line_index += 1

        flush_paragraph()
        return blocks

    @staticmethod
    def _split_table_row(line: str) -> list[str]:
        normalized = line.strip().strip("|")
        return [cell.strip() for cell in normalized.split("|")]

    @staticmethod
    def _looks_like_new_block(stripped_line: str) -> bool:
        return bool(
            HEADING_RE.match(stripped_line)
            or BULLET_RE.match(stripped_line)
            or ORDERED_RE.match(stripped_line)
            or HORIZONTAL_RULE_RE.match(stripped_line)
            or stripped_line.startswith("|")
        )

    def _parse_inline_runs(self, text: str) -> list[InlineRun]:
        runs: list[InlineRun] = []
        for fragment in INLINE_TOKEN_RE.split(text):
            if not fragment:
                continue
            if fragment.startswith("**") and fragment.endswith("**") and len(fragment) > 4:
                runs.append(InlineRun(text=fragment[2:-2], bold=True))
                continue
            if fragment.startswith("`") and fragment.endswith("`") and len(fragment) > 2:
                runs.append(InlineRun(text=fragment[1:-1], code=True))
                continue
            runs.append(InlineRun(text=fragment))
        return runs

    def _build_docx(
        self,
        target_path: Path,
        blocks: list[MarkdownBlock],
        report_markdown: str,
        trace_id: str,
    ) -> None:
        self._ensure_docx_dependencies()
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

        self._configure_docx_base_style(document)
        document.core_properties.title = "交通事故分析报告"
        document.core_properties.subject = trace_id

        title_written = False
        for block in blocks:
            if block.kind == "heading":
                paragraph = document.add_paragraph()
                self._render_docx_heading(paragraph, block, is_title=not title_written and block.level == 1)
                title_written = title_written or block.level == 1
                continue

            if block.kind == "paragraph":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(10)
                paragraph.paragraph_format.line_spacing = 1.4
                self._append_runs_to_docx(paragraph, block.runs)
                continue

            if block.kind == "list":
                for index, item in enumerate(block.items, start=1):
                    paragraph = document.add_paragraph(
                        style="List Number" if block.ordered else "List Bullet"
                    )
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.line_spacing = 1.35
                    if block.ordered:
                        paragraph.text = ""
                    self._append_runs_to_docx(paragraph, item)
                    if block.ordered:
                        paragraph.style = document.styles["List Number"]
                continue

            if block.kind == "table":
                self._render_docx_table(document, block.rows)
                document.add_paragraph()
                continue

            if block.kind == "hr":
                divider = document.add_paragraph()
                divider.paragraph_format.space_before = Pt(6)
                divider.paragraph_format.space_after = Pt(10)
                run = divider.add_run(" ")
                border = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "CBD5E1")
                border.append(bottom)
                divider._p.get_or_add_pPr().append(border)
                self._set_docx_run_font(run)

        if not blocks:
            document.add_paragraph(report_markdown)

        target_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(target_path)

    def _configure_docx_base_style(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Microsoft YaHei"
        normal_style.font.size = Pt(11)
        style_r_pr = normal_style._element.get_or_add_rPr()
        style_r_fonts = style_r_pr.rFonts
        if style_r_fonts is None:
            style_r_fonts = OxmlElement("w:rFonts")
            style_r_pr.append(style_r_fonts)
        style_r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def _render_docx_heading(self, paragraph, block: MarkdownBlock, is_title: bool) -> None:  # noqa: ANN001
        paragraph.paragraph_format.space_before = Pt(14 if block.level > 1 else 6)
        paragraph.paragraph_format.space_after = Pt(8 if block.level > 1 else 14)
        paragraph.paragraph_format.line_spacing = 1.2
        if is_title:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size_map = {
            1: Pt(20 if is_title else 16),
            2: Pt(15),
            3: Pt(13),
            4: Pt(12),
            5: Pt(11),
            6: Pt(11),
        }
        color_map = {
            1: RGBColor(15, 23, 42),
            2: RGBColor(30, 64, 175),
            3: RGBColor(51, 65, 85),
        }
        for inline_run in block.runs:
            run = paragraph.add_run(inline_run.text)
            run.bold = True
            run.font.size = size_map.get(block.level, Pt(11))
            run.font.color.rgb = color_map.get(block.level, RGBColor(15, 23, 42))
            self._set_docx_run_font(run, code=inline_run.code)

    def _append_runs_to_docx(self, paragraph, runs: list[InlineRun]) -> None:  # noqa: ANN001
        for inline_run in runs:
            run = paragraph.add_run(inline_run.text)
            run.bold = inline_run.bold
            self._set_docx_run_font(run, code=inline_run.code)
            if inline_run.code:
                run.font.color.rgb = RGBColor(185, 28, 28)

    def _set_docx_run_font(self, run, code: bool = False) -> None:  # noqa: ANN001
        font_name = "Consolas" if code else "Microsoft YaHei"
        east_asia_font = "Consolas" if code else "Microsoft YaHei"
        run.font.name = font_name
        run_r_pr = run._element.get_or_add_rPr()
        run_r_fonts = run_r_pr.rFonts
        if run_r_fonts is None:
            run_r_fonts = OxmlElement("w:rFonts")
            run_r_pr.append(run_r_fonts)
        run_r_fonts.set(qn("w:eastAsia"), east_asia_font)

    def _render_docx_table(self, document: Document, rows: list[list[list[InlineRun]]]) -> None:
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            for col_index in range(column_count):
                cell = table.cell(row_index, col_index)
                paragraph = cell.paragraphs[0]
                paragraph.text = ""
                if col_index < len(row):
                    self._append_runs_to_docx(paragraph, row[col_index])
                if row_index == 0:
                    self._set_docx_cell_shading(cell, "E2E8F0")
                    for run in paragraph.runs:
                        run.bold = True

    @staticmethod
    def _set_docx_cell_shading(cell, fill: str) -> None:  # noqa: ANN001
        cell_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_pr.append(shading)

    def _build_pdf(
        self,
        target_path: Path,
        blocks: list[MarkdownBlock],
        trace_id: str,
        cover_config: ResolvedPdfCoverConfig,
    ) -> None:
        self._ensure_pdf_dependencies()
        self._ensure_pdf_font_registered()
        target_path.parent.mkdir(parents=True, exist_ok=True)
        document = _StyledPdfDocTemplate(
            str(target_path),
            title=cover_config.title,
            author=cover_config.compiled_by,
            subject=trace_id,
        )
        styles = self._build_pdf_styles()
        body_blocks = self._strip_pdf_cover_title_block(blocks)
        story = []
        story.extend(self._build_pdf_cover_story(styles, cover_config))
        story.append(NextPageTemplate("Body"))
        story.append(PageBreak())
        story.extend(self._build_pdf_toc_story(styles, body_blocks))
        story.append(PageBreak())
        story.extend(self._build_pdf_body_story(styles, body_blocks))

        document.multiBuild(story)

    def _resolve_pdf_cover_config(
        self,
        body_blocks: list[MarkdownBlock],
        trace_id: str,
        options: PdfCoverOptions | None,
    ) -> ResolvedPdfCoverConfig:
        normalized_options = options or PdfCoverOptions()
        title = self._normalize_pdf_cover_text(
            normalized_options.title,
            fallback=PDF_COVER_TITLE,
            field_label="封面标题",
            max_length=48,
        )
        subtitle = self._normalize_pdf_cover_text(
            normalized_options.subtitle,
            fallback=self._resolve_pdf_cover_subtitle(body_blocks),
            field_label="封面副标题",
            max_length=64,
        )
        compiled_by = self._normalize_pdf_cover_text(
            normalized_options.compiled_by,
            fallback=PDF_COMPILED_BY,
            field_label="编制人",
            max_length=48,
        )
        date_mode = normalized_options.date_mode or "today"
        if date_mode == "custom":
            date_text = self._normalize_pdf_cover_text(
                normalized_options.date_text,
                fallback="",
                field_label="封面日期",
                max_length=32,
                allow_empty=False,
            )
        elif date_mode == "hide":
            date_text = None
        else:
            date_text = datetime.now().strftime("%Y年%m月%d日")

        return ResolvedPdfCoverConfig(
            title=title,
            subtitle=subtitle,
            compiled_by=compiled_by,
            date_text=date_text,
            report_number=trace_id,
        )

    def _is_pdf_export_current(
        self,
        target_path: Path,
        source_mtime_ns: int,
        template_mtime_ns: int,
        cover_config: ResolvedPdfCoverConfig,
    ) -> bool:
        meta_path = self._get_pdf_export_meta_path(target_path)
        if not target_path.exists() or not meta_path.exists():
            return False
        if target_path.stat().st_mtime_ns < max(source_mtime_ns, template_mtime_ns):
            return False
        try:
            meta_payload = json.loads(meta_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return False

        return meta_payload == {
            "report_markdown_mtime_ns": source_mtime_ns,
            "template_mtime_ns": template_mtime_ns,
            "cover": cover_config.to_cache_payload(),
        }

    def _write_pdf_export_meta(
        self,
        target_path: Path,
        source_mtime_ns: int,
        template_mtime_ns: int,
        cover_config: ResolvedPdfCoverConfig,
    ) -> None:
        meta_path = self._get_pdf_export_meta_path(target_path)
        meta_payload = {
            "report_markdown_mtime_ns": source_mtime_ns,
            "template_mtime_ns": template_mtime_ns,
            "cover": cover_config.to_cache_payload(),
        }
        meta_path.write_text(json.dumps(meta_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    @staticmethod
    def _get_pdf_export_meta_path(target_path: Path) -> Path:
        return target_path.with_suffix(f"{target_path.suffix}.meta.json")

    @staticmethod
    def _normalize_pdf_cover_text(
        value: str | None,
        fallback: str,
        field_label: str,
        max_length: int,
        allow_empty: bool = False,
    ) -> str:
        candidate = str(value or "").strip()
        if not candidate:
            candidate = fallback.strip()
        if not candidate and not allow_empty:
            raise InputValidationError(f"{field_label}不能为空。")
        if len(candidate) > max_length:
            raise InputValidationError(f"{field_label}不能超过 {max_length} 个字符。")
        return candidate

    @classmethod
    def _ensure_pdf_font_registered(cls) -> None:
        cls._ensure_pdf_dependencies()
        if cls._pdf_font_registered:
            return
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        cls._pdf_font_registered = True

    @staticmethod
    def _ensure_docx_dependencies() -> None:
        if DOCX_IMPORT_ERROR is None:
            return
        raise DependencyUnavailableError(
            "Word 导出依赖未安装，请在项目虚拟环境中执行 "
            "`python -m pip install -r backend/requirements.txt` 后重试。"
        )

    @staticmethod
    def _ensure_pdf_dependencies() -> None:
        if PDF_IMPORT_ERROR is None:
            return
        raise DependencyUnavailableError(
            "PDF 导出依赖未安装，请在项目虚拟环境中执行 "
            "`python -m pip install -r backend/requirements.txt` 后重试。"
        )

    @staticmethod
    def _build_pdf_styles() -> dict[str, ParagraphStyle]:
        stylesheet = getSampleStyleSheet()
        return {
            "CoverBrand": ParagraphStyle(
                "CoverBrand",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=11.2,
                leading=15,
                textColor=colors.HexColor(PDF_COLOR_ACCENT),
                wordWrap="CJK",
                spaceAfter=0.18 * cm,
            ),
            "CoverTitle": ParagraphStyle(
                "CoverTitle",
                parent=stylesheet["Title"],
                fontName="STSong-Light",
                fontSize=28,
                leading=37,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                spaceAfter=0.18 * cm,
            ),
            "CoverSubtitle": ParagraphStyle(
                "CoverSubtitle",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=12.3,
                leading=18,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
                spaceAfter=0.24 * cm,
            ),
            "CoverMetaLabel": ParagraphStyle(
                "CoverMetaLabel",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=8.8,
                leading=12,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
            ),
            "CoverMetaValue": ParagraphStyle(
                "CoverMetaValue",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=11.4,
                leading=16,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
            ),
            "CoverMetaHint": ParagraphStyle(
                "CoverMetaHint",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=9.2,
                leading=14,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
            ),
            "TOCTitle": ParagraphStyle(
                "TOCTitle",
                parent=stylesheet["Heading1"],
                fontName="STSong-Light",
                fontSize=19,
                leading=24,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                spaceAfter=0.18 * cm,
            ),
            "TOCLead": ParagraphStyle(
                "TOCLead",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=10,
                leading=14,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
                spaceAfter=0.42 * cm,
            ),
            "TOCEmpty": ParagraphStyle(
                "TOCEmpty",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=10.5,
                leading=16,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
                spaceAfter=0.2 * cm,
            ),
            "H1CN": ParagraphStyle(
                "H1CN",
                parent=stylesheet["Heading1"],
                fontName="STSong-Light",
                fontSize=17,
                leading=24,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                spaceBefore=0.1 * cm,
                spaceAfter=0.25 * cm,
            ),
            "H2CN": ParagraphStyle(
                "H2CN",
                parent=stylesheet["Heading2"],
                fontName="STSong-Light",
                fontSize=15,
                leading=20,
                textColor=colors.HexColor(PDF_COLOR_HEADER),
                wordWrap="CJK",
                spaceBefore=10,
                spaceAfter=8,
            ),
            "H3CN": ParagraphStyle(
                "H3CN",
                parent=stylesheet["Heading3"],
                fontName="STSong-Light",
                fontSize=12.5,
                leading=18,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
                spaceBefore=8,
                spaceAfter=6,
            ),
            "BodyCN": ParagraphStyle(
                "BodyCN",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=10.5,
                leading=17,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                spaceAfter=8,
            ),
            "ListCN": ParagraphStyle(
                "ListCN",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=10.5,
                leading=17,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                spaceAfter=4,
                leftIndent=10,
            ),
            "TableCellCN": ParagraphStyle(
                "TableCellCN",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=9.6,
                leading=13.5,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
            ),
            "TOCLevel1": ParagraphStyle(
                "TOCLevel1",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=11,
                leading=17,
                textColor=colors.HexColor(PDF_COLOR_INK),
                wordWrap="CJK",
                leftIndent=0,
                firstLineIndent=0,
                spaceBefore=2,
            ),
            "TOCLevel2": ParagraphStyle(
                "TOCLevel2",
                parent=stylesheet["BodyText"],
                fontName="STSong-Light",
                fontSize=10,
                leading=15,
                textColor=colors.HexColor(PDF_COLOR_MUTED),
                wordWrap="CJK",
                leftIndent=18,
                firstLineIndent=0,
                spaceBefore=1,
            ),
        }

    def _build_pdf_cover_story(
        self,
        styles: dict[str, ParagraphStyle],
        cover_config: ResolvedPdfCoverConfig,
    ) -> list:
        meta_columns: list[list] = [
            [
                Paragraph("编制人", styles["CoverMetaLabel"]),
                Paragraph(cover_config.compiled_by, styles["CoverMetaValue"]),
            ],
        ]
        if cover_config.date_text:
            meta_columns.append(
                [
                    Paragraph("编制日期", styles["CoverMetaLabel"]),
                    Paragraph(cover_config.date_text, styles["CoverMetaValue"]),
                ]
            )
        meta_columns.append(
            [
                Paragraph("报告编号", styles["CoverMetaLabel"]),
                Paragraph(cover_config.report_number, styles["CoverMetaValue"]),
            ]
        )
        column_width = 17 * cm / len(meta_columns)
        meta_table = Table([meta_columns], colWidths=[column_width] * len(meta_columns))
        meta_table.setStyle(
            TableStyle(
                [
                    ("LINEABOVE", (0, 0), (-1, 0), 1.0, colors.HexColor(PDF_COLOR_ACCENT)),
                    ("LINEBELOW", (0, 0), (-1, 0), 0.55, colors.HexColor(PDF_COLOR_LINE)),
                    ("LINEBEFORE", (1, 0), (-1, 0), 0.45, colors.HexColor(PDF_COLOR_LINE)),
                    ("VALIGN", (0, 0), (-1, 0), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, 0), 10),
                    ("RIGHTPADDING", (0, 0), (-1, 0), 10),
                    ("TOPPADDING", (0, 0), (-1, 0), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ]
            )
        )

        return [
            Spacer(1, 3.45 * cm),
            Paragraph("锐鉴安途道路交通事故分析系统", styles["CoverBrand"]),
            Paragraph(cover_config.title, styles["CoverTitle"]),
            Paragraph(cover_config.subtitle, styles["CoverSubtitle"]),
            Paragraph("用于事故事实梳理、责任分析与研判论述的归档版文书输出。", styles["CoverMetaHint"]),
            Spacer(1, 5.6 * cm),
            meta_table,
        ]

    def _build_pdf_toc_story(self, styles: dict[str, ParagraphStyle], body_blocks: list[MarkdownBlock]) -> list:
        flowables = [
            Paragraph("目录", styles["TOCTitle"]),
            Paragraph("以下目录根据正文标题自动生成，便于快速定位分析章节。", styles["TOCLead"]),
        ]
        if not self._has_pdf_toc_entries(body_blocks):
            flowables.append(Paragraph("正文未检测到可编目的标题，将直接显示报告正文。", styles["TOCEmpty"]))
            return flowables

        toc = TableOfContents()
        toc.levelStyles = [styles["TOCLevel1"], styles["TOCLevel2"]]
        toc.dotsMinLevel = 0
        flowables.append(toc)
        return flowables

    def _build_pdf_body_story(self, styles: dict[str, ParagraphStyle], body_blocks: list[MarkdownBlock]) -> list:
        flowables = []

        for block in body_blocks:
            if block.kind == "heading":
                style_name = "H1CN" if block.level == 1 else "H2CN" if block.level == 2 else "H3CN"
                paragraph = Paragraph(self._runs_to_pdf_markup(block.runs), styles[style_name])
                self._mark_pdf_heading(paragraph, block)
                flowables.append(paragraph)
                continue

            if block.kind == "paragraph":
                flowables.append(Paragraph(self._runs_to_pdf_markup(block.runs), styles["BodyCN"]))
                continue

            if block.kind == "list":
                for index, item in enumerate(block.items, start=1):
                    marker = f"{index}. " if block.ordered else "• "
                    flowables.append(Paragraph(f"{escape(marker)}{self._runs_to_pdf_markup(item)}", styles["ListCN"]))
                continue

            if block.kind == "table":
                if not block.rows:
                    continue
                table_data = [
                    [Paragraph(self._runs_to_pdf_markup(cell), styles["TableCellCN"]) for cell in row]
                    for row in block.rows
                ]
                table = Table(table_data, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                            ("FONTSIZE", (0, 0), (-1, -1), 10),
                            ("LEADING", (0, 0), (-1, -1), 14),
                            ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1E293B")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E2E8F0")),
                            ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#CBD5E1")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 8),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    )
                )
                flowables.append(table)
                flowables.append(Spacer(1, 0.18 * cm))
                continue

            if block.kind == "hr":
                flowables.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#CBD5E1")))
                flowables.append(Spacer(1, 0.18 * cm))

        if not flowables:
            flowables.append(Paragraph("报告内容为空。", styles["BodyCN"]))

        return flowables

    def _resolve_pdf_cover_subtitle(self, body_blocks: list[MarkdownBlock]) -> str:
        first_heading = self._find_first_heading_text(body_blocks)
        if first_heading and first_heading not in PDF_GENERIC_BODY_TITLES:
            return first_heading
        return PDF_COVER_SUBTITLE

    def _strip_pdf_cover_title_block(self, blocks: list[MarkdownBlock]) -> list[MarkdownBlock]:
        if not blocks:
            return []
        first_block = blocks[0]
        if first_block.kind != "heading":
            return blocks
        first_heading = self._plain_text_from_runs(first_block.runs)
        if first_heading in PDF_GENERIC_BODY_TITLES:
            return blocks[1:]
        return blocks

    def _has_pdf_toc_entries(self, blocks: list[MarkdownBlock]) -> bool:
        return any(block.kind == "heading" and self._map_pdf_toc_level(block.level) is not None for block in blocks)

    def _mark_pdf_heading(self, paragraph, block: MarkdownBlock) -> None:  # noqa: ANN001
        toc_level = self._map_pdf_toc_level(block.level)
        if toc_level is None:
            return
        paragraph._toc_level = toc_level
        paragraph._plain_text = self._plain_text_from_runs(block.runs)

    @staticmethod
    def _map_pdf_toc_level(heading_level: int) -> int | None:
        if heading_level <= 2:
            return 0
        if heading_level <= 4:
            return 1
        return None

    @staticmethod
    def _find_first_heading_text(blocks: list[MarkdownBlock]) -> str:
        for block in blocks:
            if block.kind == "heading":
                return ReportExportService._plain_text_from_runs(block.runs)
        return ""

    @staticmethod
    def _plain_text_from_runs(runs: list[InlineRun]) -> str:
        return "".join(run.text for run in runs).strip()

    def _runs_to_pdf_markup(self, runs: list[InlineRun]) -> str:
        fragments: list[str] = []
        for inline_run in runs:
            text = escape(inline_run.text).replace("\n", "<br/>")
            if inline_run.bold:
                fragments.append(f"<b>{text}</b>")
                continue
            if inline_run.code:
                fragments.append(f'<font color="#B91C1C">{text}</font>')
                continue
            fragments.append(text)
        return "".join(fragments) or "&nbsp;"
